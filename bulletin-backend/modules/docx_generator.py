"""
DOCX Bulletin Generator - Episcopal / BCP Style
Generates formatted Word documents for Sunday bulletins.
Uses python-docx for document creation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def _safe_get(data, key, default=""):
    """Safely get a value from data dict."""
    return data.get(key, default) or default


def setup_styles(doc):
    """Configure document styles for BCP bulletin formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Garamond"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(3)

    # Configure heading styles
    for level, size in [(0, 22), (1, 16), (2, 13)]:
        heading_style = doc.styles[f"Heading {level + 1}" if level > 0 else "Title"]
        if hasattr(heading_style, "font"):
            heading_style.font.name = "Garamond"
            heading_style.font.size = Pt(size)
            heading_style.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Dark red


def add_header(doc, data):
    """Add parish name, service info header."""
    # Parish name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(_safe_get(data, "parish_name", "Episcopal Church"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("" + "\u2500" * 40 + "")
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Service type
    service_type = _safe_get(data, "service_type", "Holy Eucharist Rite II")
    stype = doc.add_paragraph()
    stype.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stype.add_run(service_type)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Garamond"

    # Date and time
    date_str = _safe_get(data, "service_date")
    time_str = _safe_get(data, "service_time")
    season = _safe_get(data, "liturgical_season")

    info_parts = []
    if date_str:
        info_parts.append(date_str)
    if time_str:
        info_parts.append(time_str)
    if season:
        info_parts.append(season)

    if info_parts:
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(" \u2022 ".join(info_parts))
        run.font.size = Pt(12)
        run.font.name = "Garamond"
        run.italic = True

    # Spacer
    doc.add_paragraph()


def add_hymn(doc, label, number, data):
    """Add a hymn entry with number, title, and tune."""
    if not number:
        return

    title = _safe_get(data, f"{label}_title", "")
    tune = _safe_get(data, f"{label}_tune", "")

    p = doc.add_paragraph()
    # Label
    run = p.add_run(f"{label.replace('_', ' ').title()}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Garamond"

    # Hymn info
    hymn_text = f"Hymn {number}"
    if title:
        hymn_text += f"  \u2013  {title}"
    if tune:
        hymn_text += f"  ({tune})"

    run = p.add_run(hymn_text)
    run.font.size = Pt(11)
    run.font.name = "Garamond"


def add_music_section(doc, data):
    """Add opening hymn and liturgical music settings."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("THE WORD OF GOD")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    # Opening hymn
    add_hymn(doc, "opening_hymn", _safe_get(data, "opening_hymn_number"), data)

    # Gloria / liturgical setting
    gloria = _safe_get(data, "gloria_number")
    if gloria:
        p = doc.add_paragraph()
        run = p.add_run("Gloria in excelsis: ")
        run.bold = True
        run.font.name = "Garamond"
        run = p.add_run(gloria)
        run.font.name = "Garamond"


def add_scripture_section(doc, data):
    """Add scripture readings in a formatted table."""
    # Collect readings
    readings = []
    first = _safe_get(data, "first_lesson_citation")
    if first:
        readings.append(("First Lesson", first))

    psalm = _safe_get(data, "psalm_number")
    if psalm:
        readings.append(("Psalm", psalm))

    second = _safe_get(data, "second_lesson_citation")
    if second:
        readings.append(("Second Lesson", second))

    # Sequence hymn between second lesson and gospel
    add_hymn(doc, "sequence_hymn", _safe_get(data, "sequence_hymn_number"), data)

    gospel = _safe_get(data, "gospel_citation")
    if gospel:
        readings.append(("The Holy Gospel", gospel))

    if not readings:
        return

    heading = doc.add_paragraph()
    run = heading.add_run("The Lessons")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    # Table for readings
    table = doc.add_table(rows=len(readings), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, citation) in enumerate(readings):
        row = table.rows[i]
        # Label cell
        cell_label = row.cells[0]
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Garamond"

        # Citation cell
        cell_cite = row.cells[1]
        p = cell_cite.paragraphs[0]
        run = p.add_run(citation)
        run.font.size = Pt(11)
        run.font.name = "Garamond"
        if label == "The Holy Gospel":
            run.italic = True

    doc.add_paragraph()


def add_sermon_section(doc, data):
    """Add sermon information."""
    sermon_title = _safe_get(data, "sermon_title")
    preacher = _safe_get(data, "preacher_name")

    if not sermon_title and not preacher:
        p = doc.add_paragraph()
        run = p.add_run("The Sermon")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        return

    heading = doc.add_paragraph()
    run = heading.add_run("The Sermon")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    if sermon_title:
        p = doc.add_paragraph()
        run = p.add_run(f'"{sermon_title}"')
        run.italic = True
        run.font.name = "Garamond"

    if preacher:
        p = doc.add_paragraph()
        run = p.add_run(f"The {preacher}")
        run.font.name = "Garamond"


def add_prayers_section(doc, data):
    """Add Prayers of the People placeholder."""
    heading = doc.add_paragraph()
    run = heading.add_run("The Prayers of the People")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("The Nicene Creed")
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("BCP p. 358")
    run.font.name = "Garamond"
    run.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Prayers of the People")
    run.bold = True
    run.font.name = "Garamond"

    p = doc.add_paragraph()
    run = p.add_run("Confession of Sin")
    run.bold = True
    run.font.name = "Garamond"

    p = doc.add_paragraph()
    run = p.add_run("BCP p. 360")
    run.font.name = "Garamond"
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run("The Peace")
    run.bold = True
    run.font.name = "Garamond"

    doc.add_paragraph()


def add_communion_section(doc, data):
    """Add Holy Communion section."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("THE HOLY COMMUNION")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    # Offertory / communion hymns
    add_hymn(doc, "communion_hymn_1", _safe_get(data, "communion_hymn_1_number"), data)
    add_hymn(doc, "communion_hymn_2", _safe_get(data, "communion_hymn_2_number"), data)

    # Sanctus
    sanctus = _safe_get(data, "sanctus_number")
    if sanctus:
        p = doc.add_paragraph()
        run = p.add_run("Sanctus: ")
        run.bold = True
        run.font.name = "Garamond"
        run = p.add_run(sanctus)
        run.font.name = "Garamond"

    # Eucharistic Prayer
    p = doc.add_paragraph()
    run = p.add_run("The Great Thanksgiving")
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Eucharistic Prayer A  \u2013  BCP p. 361")
    run.font.name = "Garamond"
    run.italic = True

    # Fraction
    fraction = _safe_get(data, "fraction_number")
    if fraction:
        p = doc.add_paragraph()
        run = p.add_run("Fraction Anthem: ")
        run.bold = True
        run.font.name = "Garamond"
        run = p.add_run(fraction)
        run.font.name = "Garamond"

    # Communion
    p = doc.add_paragraph()
    run = p.add_run("The Communion of the People")
    run.bold = True
    run.font.name = "Garamond"

    p = doc.add_paragraph()
    run = p.add_run("All baptized Christians are welcome to receive Holy Communion.")
    run.font.name = "Garamond"
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()


def add_closing_section(doc, data):
    """Add closing hymn and dismissal."""
    # Post-communion prayer
    p = doc.add_paragraph()
    run = p.add_run("Post-Communion Prayer")
    run.bold = True
    run.font.name = "Garamond"

    p = doc.add_paragraph()
    run = p.add_run("BCP p. 365")
    run.font.name = "Garamond"
    run.italic = True

    # Blessing
    p = doc.add_paragraph()
    run = p.add_run("The Blessing")
    run.bold = True
    run.font.name = "Garamond"

    doc.add_paragraph()

    # Closing hymn
    add_hymn(doc, "closing_hymn", _safe_get(data, "closing_hymn_number"), data)

    # Dismissal
    p = doc.add_paragraph()
    run = p.add_run("The Dismissal")
    run.bold = True
    run.font.name = "Garamond"


def add_participants_section(doc, data):
    """Add ministers/participants table."""
    participants = []
    rector = _safe_get(data, "rector_name")
    if rector:
        participants.append(("Rector", rector))

    preacher = _safe_get(data, "preacher_name")
    if preacher:
        participants.append(("Preacher", preacher))

    music_dir = _safe_get(data, "music_director_name")
    if music_dir:
        participants.append(("Music Director", music_dir))

    organist = _safe_get(data, "organist_name")
    if organist:
        participants.append(("Organist", organist))

    if not participants:
        return

    doc.add_paragraph()

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("\u2500" * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    table = doc.add_table(rows=len(participants), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (role, name) in enumerate(participants):
        row = table.rows[i]
        cell_role = row.cells[0]
        p = cell_role.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(role)
        run.font.size = Pt(10)
        run.font.name = "Garamond"
        run.italic = True

        cell_name = row.cells[1]
        p = cell_name.paragraphs[0]
        run = p.add_run(name)
        run.font.size = Pt(10)
        run.font.name = "Garamond"


def add_footer(doc, data):
    """Add parish contact information footer."""
    address = _safe_get(data, "parish_address")
    phone = _safe_get(data, "parish_phone")
    website = _safe_get(data, "parish_website")

    if not address and not phone and not website:
        return

    doc.add_paragraph()

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("\u2500" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer_parts = [p for p in [address, phone, website] if p]
    footer_text = " \u2022 ".join(footer_parts)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(footer_text)
    run.font.size = Pt(9)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def generate_bulletin(data, output_path):
    """
    Main entry point: generate a complete Episcopal bulletin DOCX.

    Args:
        data: dict with bulletin form fields
        output_path: str path where the DOCX file should be saved

    Returns:
        str: the output path of the generated file
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(5.5)   # Half-letter width
    section.page_height = Inches(8.5)  # Half-letter height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Apply styles
    setup_styles(doc)

    # Build bulletin sections
    add_header(doc, data)
    add_music_section(doc, data)
    add_scripture_section(doc, data)
    add_sermon_section(doc, data)
    add_prayers_section(doc, data)
    add_communion_section(doc, data)
    add_closing_section(doc, data)
    add_participants_section(doc, data)
    add_footer(doc, data)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True) if os.path.dirname(output_path) else None

    doc.save(output_path)
    return output_path
